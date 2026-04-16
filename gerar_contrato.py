"""
Gerador de Contrato de Locação - ATIVUZ
Módulo reutilizável pela aplicação Flask.
"""

import json
import re
import tempfile
import zipfile
from datetime import datetime
from pathlib import Path


CAMPOS = {
    "locatario_nome":     "Nome completo do LOCATÁRIO",
    "locatario_rg":       "RG do LOCATÁRIO",
    "locatario_cpf":      "CPF do LOCATÁRIO",
    "locatario_endereco": "Endereço do LOCATÁRIO",
    "locatario_cep":      "CEP do LOCATÁRIO",
    "locatario_telefone": "Telefone do LOCATÁRIO",
    "avalista_nome":      "Nome completo do AVALISTA",
    "avalista_cpf":       "CPF do AVALISTA",
    "avalista_endereco":  "Endereço do AVALISTA",
    "avalista_telefone":  "Telefone do AVALISTA",
    "veiculo_descricao":  "Descrição do veículo",
    "veiculo_marca":      "Marca do veículo",
    "veiculo_modelo":     "Modelo do veículo",
    "veiculo_ano":        "Ano de fabricação",
    "veiculo_motor":      "Número do motor",
    "veiculo_chassi":     "Número do chassi",
    "veiculo_cor":        "Cor do veículo",
    "veiculo_placa":      "Placa do veículo",
    "contrato_inicio":    "Data de início da locação",
    "contrato_duracao":   "Duração em meses",
    "valor_semanal":      "Valor semanal da locação (R$)",
    "data_dia":           "Dia da assinatura",
    "data_mes":           "Mês da assinatura por extenso",
    "data_ano":           "Ano da assinatura",
    "testemunha1_nome":   "Nome da Testemunha 1",
    "testemunha1_rg":     "RG da Testemunha 1",
    "testemunha1_cpf":    "CPF da Testemunha 1",
    "testemunha2_nome":   "Nome da Testemunha 2",
    "testemunha2_rg":     "RG da Testemunha 2",
    "testemunha2_cpf":    "CPF da Testemunha 2",
}

# Ordem exata das 43 ocorrências de [ ] no template
ORDEM_OCORRENCIAS = [
    "locatario_nome",
    "locatario_rg",
    "locatario_cpf",
    "locatario_endereco",
    "locatario_cep",
    "locatario_telefone",
    "avalista_nome",
    "avalista_cpf",
    "avalista_endereco",
    "avalista_telefone",
    "locatario_telefone",
    "veiculo_descricao",
    "veiculo_marca",
    "veiculo_modelo",
    "veiculo_ano",
    "veiculo_motor",
    "veiculo_chassi",
    "veiculo_cor",
    "veiculo_placa",
    "contrato_inicio",
    "contrato_duracao",
    "valor_semanal",
    "locatario_nome",
    "locatario_rg",
    "locatario_cpf",
    "locatario_endereco",
    "locatario_cep",
    "locatario_telefone",
    "avalista_nome",
    "avalista_cpf",
    "avalista_endereco",
    "avalista_telefone",
    "data_dia",
    "data_mes",
    "data_ano",
    "locatario_nome",
    "locatario_cpf",
    "avalista_nome",
    "avalista_cpf",
    "testemunha1_nome",
    "testemunha1_rg",
    "testemunha1_cpf",
    "testemunha2_nome",
    "testemunha2_rg",
    "testemunha2_cpf",
]


def substituir_campos(xml_text: str, dados: dict) -> str:
    padrao = re.compile(r'\[ \]')
    contador = [0]

    def substituir(match):
        idx = contador[0]
        contador[0] += 1
        if idx < len(ORDEM_OCORRENCIAS):
            chave = ORDEM_OCORRENCIAS[idx]
            valor = dados.get(chave, "").strip()
            return valor if valor else "[ ]"
        return match.group(0)

    return padrao.sub(substituir, xml_text)


def gerar_docx(dados: dict, caminho_saida: str, template_path: str = None):
    """Descompacta o template, substitui os campos e recompacta."""
    template = Path(template_path) if template_path else Path(__file__).parent / "TEMPLATE_CODE.docx"

    if not template.exists():
        raise FileNotFoundError(f"Template não encontrado: {template}")

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        with zipfile.ZipFile(template, "r") as z:
            z.extractall(tmp)

        doc_xml = tmp / "word" / "document.xml"
        texto = doc_xml.read_text(encoding="utf-8")
        texto_modificado = substituir_campos(texto, dados)
        doc_xml.write_text(texto_modificado, encoding="utf-8")

        saida = Path(caminho_saida)
        saida.parent.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(saida, "w", zipfile.ZIP_DEFLATED) as z:
            for arquivo in tmp.rglob("*"):
                if arquivo.is_file():
                    z.write(arquivo, arquivo.relative_to(tmp))


def nome_arquivo_saida(dados: dict) -> str:
    nome = dados.get("locatario_nome", "contrato").split()[0].lower()
    hoje = datetime.now().strftime("%Y%m%d_%H%M%S")
    return f"contrato_{nome}_{hoje}.docx"


def gerar_termo_quitacao(
    devedor_nome: str,
    devedor_cpf: str,
    placa: str,
    mes_referencia_fipe: str,
    valor_fipe: float,
    percentual_fipe: float,
    meias_diarias: float,
    entrada: float,
    num_parcelas_pagas: int,
    valor_parcela_paga: float,
    num_parcelas_semanais: int,
    valor_parcela_semanal: float,
    data_primeira_parcela: str,
    data_assinatura: str,
    caminho_saida: str,
    template_path: str = None,
):
    """Gera o Termo de Quitação a partir do template TERMO_QUITACAO_TEMPLATE.docx."""
    from decimal import Decimal
    from num2words import num2words

    def _fmt(valor: float) -> str:
        """Formata valor monetário no padrão BR: 1.234,56"""
        return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def _ext(valor: float) -> str:
        """Retorna valor monetário por extenso em português via num2words currency."""
        return num2words(Decimal(str(round(valor, 2))), lang="pt_BR", to="currency")

    # Cálculos
    valor_percentual_fipe = valor_fipe * (percentual_fipe / 100)
    total_divida          = valor_percentual_fipe + meias_diarias
    total_parcelas_pagas  = num_parcelas_pagas * valor_parcela_paga
    total_pago            = entrada + total_parcelas_pagas
    saldo_devedor         = total_divida - total_pago
    ultima_parcela        = saldo_devedor - (num_parcelas_semanais * valor_parcela_semanal)

    # 21 substituições na ordem exata dos marcadores [ ] do template
    substituicoes = [
        # 0   devedor_nome
        devedor_nome,
        # 1   devedor_cpf
        devedor_cpf,
        # 2   placa
        placa,
        # 3   mes_referencia_fipe  (ex: "dezembro 2025")
        mes_referencia_fipe,
        # 4   valor_fipe formatado  (ex: "54.724,00")
        _fmt(valor_fipe),
        # 5   valor_percentual_fipe + extenso  (ex: "10.944,80 (dez mil...)")
        f"{_fmt(valor_percentual_fipe)} ({_ext(valor_percentual_fipe)})",
        # 6   meias_diarias + extenso  (ex: "1.200,00 (um mil e duzentos reais)")
        f"{_fmt(meias_diarias)} ({_ext(meias_diarias)})",
        # 7   total_divida + extenso
        f"{_fmt(total_divida)} ({_ext(total_divida)})",
        # 8   entrada + extenso  (ex: "4.000,00 (quatro mil reais)")
        f"{_fmt(entrada)} ({_ext(entrada)})",
        # 9   descricao_parcelas_pagas  (ex: "2 parcelas de R$ 345,00 (trezentos e quarenta e cinco reais)")
        f"{num_parcelas_pagas} parcelas de R$ {_fmt(valor_parcela_paga)} ({_ext(valor_parcela_paga)})",
        # 10  total_parcelas_pagas  (ex: "690,00")
        _fmt(total_parcelas_pagas),
        # 11  total_pago  (ex: "4.690,00")
        _fmt(total_pago),
        # 12  total_pago_extenso  (ex: "(quatro mil seiscentos e noventa reais)")
        f"({_ext(total_pago)})",
        # 13  saldo_devedor  (ex: "7.454,80")
        _fmt(saldo_devedor),
        # 14  saldo_devedor_extenso  (ex: "(sete mil quatrocentos e cinquenta e quatro reais e oitenta centavos)")
        f"({_ext(saldo_devedor)})",
        # 15  descricao_parcelamento  (ex: "37 parcelas semanais de R$ 200,00 (duzentos reais)")
        f"{num_parcelas_semanais} parcelas semanais de R$ {_fmt(valor_parcela_semanal)} ({_ext(valor_parcela_semanal)})",
        # 16  descricao_ultima_parcela  (ex: "1 última parcela no valor de R$ 54,80 (cinquenta e quatro reais e oitenta centavos)")
        f"1 última parcela no valor de R$ {_fmt(ultima_parcela)} ({_ext(ultima_parcela)})",
        # 17  data_primeira_parcela
        data_primeira_parcela,
        # 18  data_assinatura
        data_assinatura,
        # 19  devedor_nome (repetição na assinatura)
        devedor_nome,
        # 20  devedor_cpf (repetição na assinatura)
        devedor_cpf,
    ]

    template = Path(template_path) if template_path else Path(__file__).parent / "TERMO_QUITACAO_TEMPLATE.docx"
    if not template.exists():
        raise FileNotFoundError(f"Template não encontrado: {template}")

    padrao = re.compile(r'\[ \]')

    def _substituir(xml_text: str) -> str:
        contador = [0]

        def _sub(match):
            idx = contador[0]
            contador[0] += 1
            if idx < len(substituicoes):
                return substituicoes[idx]
            return match.group(0)

        return padrao.sub(_sub, xml_text)

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        with zipfile.ZipFile(template, "r") as z:
            z.extractall(tmp)

        doc_xml = tmp / "word" / "document.xml"
        texto = doc_xml.read_text(encoding="utf-8")
        doc_xml.write_text(_substituir(texto), encoding="utf-8")

        saida = Path(caminho_saida)
        saida.parent.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(saida, "w", zipfile.ZIP_DEFLATED) as z:
            for arquivo in tmp.rglob("*"):
                if arquivo.is_file():
                    z.write(arquivo, arquivo.relative_to(tmp))


def gerar_notificacao_avalista(
    avalista_nome: str,
    data_contrato: str,
    locatario_nome: str,
    valor_debito: float,
    caminho_saida: str,
    template_path: str = None,
):
    """Gera a Notificação ao Avalista a partir do template NOTIFICACAO_AVALISTA_TEMPLATE.docx."""
    from decimal import Decimal
    from num2words import num2words

    def _fmt(valor: float) -> str:
        return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def _ext(valor: float) -> str:
        return num2words(Decimal(str(round(valor, 2))), lang="pt_BR", to="currency")

    # 6 substituições na ordem exata dos marcadores [ ] do template
    substituicoes = [
        # 0  avalista_nome
        avalista_nome,
        # 1  data_contrato  (ex: "12/04/2026")
        data_contrato,
        # 2  locatario_nome
        locatario_nome,
        # 3  valor_debito formatado  (ex: "1.500,00")
        _fmt(valor_debito),
        # 4  valor_extenso  (ex: "mil e quinhentos reais")
        _ext(valor_debito),
        # 5  data_geracao automática  (ex: "Natal, 14/04/2026")
        datetime.now().strftime("%d/%m/%Y"),
    ]

    template = Path(template_path) if template_path else Path(__file__).parent / "NOTIFICACAO_AVALISTA_TEMPLATE.docx"
    if not template.exists():
        raise FileNotFoundError(f"Template não encontrado: {template}")

    padrao = re.compile(r'\[ \]')

    def _substituir(xml_text: str) -> str:
        contador = [0]

        def _sub(match):
            idx = contador[0]
            contador[0] += 1
            if idx < len(substituicoes):
                return substituicoes[idx]
            return match.group(0)

        return padrao.sub(_sub, xml_text)

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        with zipfile.ZipFile(template, "r") as z:
            z.extractall(tmp)

        doc_xml = tmp / "word" / "document.xml"
        texto = doc_xml.read_text(encoding="utf-8")
        doc_xml.write_text(_substituir(texto), encoding="utf-8")

        saida = Path(caminho_saida)
        saida.parent.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(saida, "w", zipfile.ZIP_DEFLATED) as z:
            for arquivo in tmp.rglob("*"):
                if arquivo.is_file():
                    z.write(arquivo, arquivo.relative_to(tmp))


def gerar_vistoria(dados, caminho_saida: str, template_path: str) -> None:
    """Preenche o template .xlsx de vistoria preservando a formatação original."""
    import openpyxl
    import unicodedata as _uni

    def _norm(s: str) -> str:
        s = _uni.normalize("NFD", str(s).lower())
        return "".join(c for c in s if _uni.category(c) != "Mn")

    template = Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Template não encontrado: {template}")

    wb = openpyxl.load_workbook(str(template))
    agora = datetime.now()

    # Campos de texto — ordenados do mais específico ao mais genérico
    campos_texto = [
        ("preenchido por",          dados.get("vis_preenchido_por", "")),
        ("hodometro entrega",       dados.get("vis_hodometro_entrega", "")),
        ("hodometro retorno",       dados.get("vis_hodometro_retorno", "")),
        ("luzes do painel",         dados.get("vis_luzes_painel", "")),
        ("danos ou avarias",        dados.get("vis_danos_internos", "")),
        ("observacoes gerais",      dados.get("vis_observacoes_gerais", "")),
        ("descricao dos sintomas",  dados.get("vis_descricao_sintomas", "")),
        ("cliente",                 dados.get("vis_cliente", "")),
        ("telefone",                dados.get("vis_telefone", "")),
        ("endereco",                dados.get("vis_endereco", "")),
        ("combustivel",             dados.get("vis_combustivel", "")),
        ("veiculo",                 dados.get("vis_veiculo", "")),
        ("chassi",                  dados.get("vis_chassi", "")),
        ("motor",                   dados.get("vis_motor", "")),
        ("placa",                   dados.get("vis_placa", "")),
        ("cor",                     dados.get("vis_cor", "")),
        ("ano",                     dados.get("vis_ano", "")),
        ("data",                    agora.strftime("%d/%m/%Y %H:%M")),
    ]

    # Acessórios — (label_norm, valor S/N/A)
    acessorios = [
        ("calotas",         dados.get("acc_calotas", "")),
        ("buzina",          dados.get("acc_buzina", "")),
        ("doc. crlv",       dados.get("acc_doc_crlv", "")),
        ("triangulo",       dados.get("acc_triangulo", "")),
        ("antena",          dados.get("acc_antena", "")),
        ("sensor de re",    dados.get("acc_sensor_re", "")),
        ("som",             dados.get("acc_som", "")),
        ("tapetes",         dados.get("acc_tapetes", "")),
        ("limpadores",      dados.get("acc_limpadores", "")),
        ("chave de roda",   dados.get("acc_chave_roda", "")),
        ("vidros",          dados.get("acc_vidros", "")),
        ("oleo do motor",   dados.get("acc_oleo_motor", "")),
        ("alarme",          dados.get("acc_alarme", "")),
        ("lampadas",        dados.get("acc_lampadas", "")),
        ("macaco",          dados.get("acc_macaco", "")),
        ("estepe",          dados.get("acc_estepe", "")),
        ("gnv",             dados.get("acc_gnv", "")),
        ("agua",            dados.get("acc_agua", "")),
        ("borracha psg d",  dados.get("acc_borracha_psg_d", "")),
        ("borracha mtr d",  dados.get("acc_borracha_mtr_d", "")),
        ("asa urubu dd",    dados.get("acc_asa_urubu_dd", "")),
        ("asa urubu td",    dados.get("acc_asa_urubu_td", "")),
        ("tapete de mala",  dados.get("acc_tapete_mala", "")),
        ("tampa paraxq",    dados.get("acc_tampa_paraxq", "")),
        ("borracha psg t",  dados.get("acc_borracha_psg_t", "")),
        ("borracha mtr t",  dados.get("acc_borracha_mtr_t", "")),
        ("asa urubu de",    dados.get("acc_asa_urubu_de", "")),
        ("asa urubu te",    dados.get("acc_asa_urubu_te", "")),
        ("bagagito",        dados.get("acc_bagagito", "")),
        ("lingueta",        dados.get("acc_lingueta", "")),
    ]

    # ── Preencher campos de texto: busca label → escreve célula à direita ────
    filled_labels: set = set()
    for ws in wb.worksheets:
        for row in ws.iter_rows():
            for cell in row:
                if cell.value is None:
                    continue
                cell_norm = _norm(str(cell.value))
                for label, valor in campos_texto:
                    if label in cell_norm and label not in filled_labels and valor:
                        target = ws.cell(row=cell.row, column=cell.column + 1)
                        if not target.value or str(target.value).strip() == "":
                            target.value = valor
                            filled_labels.add(label)
                            break

    # ── Preencher acessórios: localiza colunas S/N/A e marca itens ──────────
    for ws in wb.worksheets:
        s_col = n_col = a_col = header_row = None
        for row in ws.iter_rows():
            for cell in row:
                if str(cell.value or "").strip().upper() == "S":
                    nc = ws.cell(row=cell.row, column=cell.column + 1)
                    ac = ws.cell(row=cell.row, column=cell.column + 2)
                    if (str(nc.value or "").strip().upper() == "N"
                            and str(ac.value or "").strip().upper() == "A"):
                        header_row = cell.row
                        s_col = cell.column
                        n_col = cell.column + 1
                        a_col = cell.column + 2
                        break
            if header_row:
                break

        if not header_row:
            continue

        col_map = {"S": s_col, "N": n_col, "A": a_col}

        for item_label, valor in acessorios:
            if not valor or valor.upper() not in col_map:
                continue
            for row in ws.iter_rows(min_row=header_row + 1):
                found = False
                for cell in row:
                    if cell.value and item_label in _norm(str(cell.value)):
                        ws.cell(row=cell.row, column=col_map[valor.upper()], value="X")
                        found = True
                        break
                if found:
                    break

    saida = Path(caminho_saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    wb.save(str(saida))


def gerar_notificacao_inadimplente(
    locatario_nome: str,
    data_contrato: str,
    valor_debito: float,
    caminho_saida: str,
    template_path: str = None,
):
    """Gera a Notificação ao Inadimplente a partir do template NOTIFICACAO_INADIMPLENTE_TEMPLATE.docx."""
    from decimal import Decimal
    from num2words import num2words

    def _fmt(valor: float) -> str:
        return f"{valor:,.2f}".replace(",", "X").replace(".", ",").replace("X", ".")

    def _ext(valor: float) -> str:
        return num2words(Decimal(str(round(valor, 2))), lang="pt_BR", to="currency")

    # 5 substituições na ordem exata dos marcadores [ ] do template
    substituicoes = [
        # 0  locatario_nome
        locatario_nome,
        # 1  data_contrato  (ex: "12/04/2026")
        data_contrato,
        # 2  valor_debito formatado  (ex: "1.500,00")
        _fmt(valor_debito),
        # 3  valor_extenso  (ex: "mil e quinhentos reais")
        _ext(valor_debito),
        # 4  data_geracao automática  (ex: "Natal, 14/04/2026")
        datetime.now().strftime("%d/%m/%Y"),
    ]

    template = Path(template_path) if template_path else Path(__file__).parent / "NOTIFICACAO_INADIMPLENTE_TEMPLATE.docx"
    if not template.exists():
        raise FileNotFoundError(f"Template não encontrado: {template}")

    padrao = re.compile(r'\[ \]')

    def _substituir(xml_text: str) -> str:
        contador = [0]

        def _sub(match):
            idx = contador[0]
            contador[0] += 1
            if idx < len(substituicoes):
                return substituicoes[idx]
            return match.group(0)

        return padrao.sub(_sub, xml_text)

    with tempfile.TemporaryDirectory() as tmp:
        tmp = Path(tmp)

        with zipfile.ZipFile(template, "r") as z:
            z.extractall(tmp)

        doc_xml = tmp / "word" / "document.xml"
        texto = doc_xml.read_text(encoding="utf-8")
        doc_xml.write_text(_substituir(texto), encoding="utf-8")

        saida = Path(caminho_saida)
        saida.parent.mkdir(parents=True, exist_ok=True)

        with zipfile.ZipFile(saida, "w", zipfile.ZIP_DEFLATED) as z:
            for arquivo in tmp.rglob("*"):
                if arquivo.is_file():
                    z.write(arquivo, arquivo.relative_to(tmp))

def gerar_vistoria_entrega(dados, fotos: list, caminho_saida: str, template_path: str) -> None:
    """Preenche o template .docx de Vistoria de Entrega com base na estrutura real do arquivo."""
    import re
    from docx import Document
    from docx.shared import Inches
    from pathlib import Path as _Path

    template = _Path(template_path)
    if not template.exists():
        raise FileNotFoundError(f"Template nao encontrado: {template}")

    doc = Document(str(template))

    WNS = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'

    def _safe_text(s: str) -> str:
        """Remove caracteres de controle invalidos em XML."""
        return re.sub(r'[\x00-\x08\x0b\x0c\x0e-\x1f]', '', str(s))

    def _xml_replace_placeholder(element, new_value: str, pattern: str = r'\[\s*\]', count: int = 1) -> bool:
        """Substitui pattern nos elementos <w:t> filhos de element. Retorna True se substituiu."""
        replaced = 0
        for t in element.iter(f'{{{WNS}}}t'):
            if not t.text:
                continue
            new_t, n = re.subn(pattern, _safe_text(new_value), t.text, count=count - replaced)
            if n:
                t.text = new_t
                replaced += n
                if replaced >= count:
                    return True
        return replaced > 0

    def _cell_fill(cell, value: str) -> None:
        """Substitui o primeiro [ ] na celula via XML."""
        if re.search(r'\[\s+\]', cell.text):
            _xml_replace_placeholder(cell._element, value, r'\[\s+\]')

    def _inline_replace(para, value: str) -> bool:
        """Substitui a primeira ocorrencia de [ ] no paragrafo via XML."""
        if re.search(r'\[\s+\]', para.text):
            return _xml_replace_placeholder(para._element, value, r'\[\s+\]')
        return False

    tbl0 = doc.tables[0]   # campos de texto (7 linhas x 4 colunas)
    tbl1 = doc.tables[1]   # checklist acessorios (11 linhas x 24 colunas)
    tbl2 = doc.tables[2]   # assinaturas (1 linha x 2 colunas)

    # ── Tabela 0: campos de texto ─────────────────────────────────────────────
    # Estrutura: C0=label, C1=valor, C2=label, C3=valor
    mapa_t0 = [
        (0, 1, dados.get("cliente", "")),
        (0, 3, dados.get("preenchido_por", "")),
        (1, 1, dados.get("endereco", "")),
        (1, 3, dados.get("tel", "")),
        (2, 1, dados.get("chassi", "")),
        (2, 3, dados.get("motor", "")),
        (3, 1, dados.get("veiculo", "")),
        (3, 3, dados.get("placa", "")),
        (4, 1, dados.get("ano", "")),
        (4, 3, dados.get("cor", "")),
        (5, 1, dados.get("hodometro_entrega", "")),
        (5, 3, dados.get("hodometro_retorno", "")),
        (6, 3, dados.get("data", "")),
    ]
    for r, c, valor in mapa_t0:
        if valor:
            _cell_fill(tbl0.rows[r].cells[c], valor)

    # ── Combustivel: R6 C1 = "[  ] MT [  ] 6/8 [  ] TC" ────────────────────
    # O texto pode estar num unico w:t ou fragmentado em varios.
    # Estrategia 1: regex direta no elemento (texto unido).
    # Estrategia 2: procura [  ] antes do w:t que contem o combustivel.
    combustivel = dados.get("combustivel", "").upper()
    if combustivel in ("MT", "6/8", "TC"):
        cell_c = tbl0.rows[6].cells[1]
        pat = r'\[\s+\](\s*' + re.escape(combustivel) + r')'
        filled = False
        for t in cell_c._element.iter(f'{{{WNS}}}t'):
            if t.text and re.search(pat, t.text):
                t.text = re.sub(pat, lambda m: '[x]' + m.group(1), t.text, count=1)
                filled = True
                break
        if not filled:
            # Texto fragmentado: acha w:t com combustivel e marca o [  ] anterior
            wt_list = [t for t in cell_c._element.iter(f'{{{WNS}}}t')]
            for idx, t in enumerate(wt_list):
                if t.text and combustivel in t.text.upper():
                    for prev in reversed(wt_list[:idx]):
                        if prev.text and re.search(r'\[\s+\]', prev.text):
                            prev.text = re.sub(r'\[\s+\]', '[x]', prev.text, count=1)
                            break
                    break

    # ── Checklist: 30 itens — mapeamento direto (chk_row, S_col, N_col, A_col, chave)
    # Estrutura VISTORIA_TESTE_1: 15 linhas x 22 colunas
    #   linhas 0,3,6,9,12: nomes dos itens
    #   linhas 1,4,7,10,13: labels S/N/A
    #   linhas 2,5,8,11,14: celulas de checkbox [    ]
    checklist = [
        # chk_row, S, N, A,  chave
        (2,  0,  1,  2,  "acc_calotas"),
        (2,  3,  5,  7,  "acc_buzina"),
        (2,  8,  9,  10, "acc_doc_crlv"),
        (2,  11, 12, 13, "acc_triangulo"),
        (2,  14, 15, 16, "acc_antena"),
        (2,  17, 19, 20, "acc_sensor_re"),
        (5,  0,  1,  2,  "acc_som"),
        (5,  3,  4,  6,  "acc_tapetes"),
        (5,  8,  9,  10, "acc_limpadores"),
        (5,  11, 12, 13, "acc_chave_roda"),
        (5,  14, 15, 16, "acc_vidros"),
        (5,  17, 19, 20, "acc_oleo_motor"),
        (8,  0,  1,  2,  "acc_alarme"),
        (8,  3,  4,  6,  "acc_lampadas"),
        (8,  8,  9,  10, "acc_macaco"),
        (8,  11, 12, 13, "acc_estepe"),
        (8,  14, 15, 16, "acc_gnv"),
        (8,  17, 19, 20, "acc_agua"),
        (11, 0,  1,  2,  "acc_borracha_psg_d"),
        (11, 3,  4,  6,  "acc_borr_mtr"),
        (11, 8,  9,  10, "acc_asa_urubu_dd"),
        (11, 11, 12, 13, "acc_asa_urub_td"),
        (11, 14, 15, 16, "acc_tapete_mala"),
        (11, 17, 19, 20, "acc_tampa_prx"),
        (14, 0,  1,  2,  "acc_borracha_psg_t"),
        (14, 3,  4,  6,  "acc_borr_mtr_t"),
        (14, 8,  9,  10, "acc_asa_urubu_de"),
        (14, 11, 12, 13, "acc_asa_urub_te"),
        (14, 14, 15, 16, "acc_bagagito"),
        (14, 17, 19, 20, "acc_linguet"),
    ]
    sna_col = {"S": 0, "N": 1, "A": 2}
    for chk_row, s_col, n_col, a_col, key in checklist:
        val = dados.get(key, "").upper()
        if val not in ("S", "N", "A"):
            continue
        col_map = {"S": s_col, "N": n_col, "A": a_col}
        target_col = col_map[val]
        cell = tbl1.rows[chk_row].cells[target_col]
        # O [    ] pode estar fragmentado em varios w:t; escreve direto na celula
        for para in cell.paragraphs:
            para.clear()
            para.add_run('[x]')
            break

    # ── Paragrafos soltos: OBSERVACOES e SINTOMAS ─────────────────────────────
    all_paras = list(doc.paragraphs)
    for i, para in enumerate(all_paras):
        heading = para.text.strip().upper()
        if "OBSERVA" in heading and i + 1 < len(all_paras):
            obs = dados.get("observacoes", "")
            if obs:
                nxt = all_paras[i + 1]
                if re.search(r'\[\s+\]', nxt.text):
                    _inline_replace(nxt, obs)
        elif "SINTOMAS" in heading and i + 1 < len(all_paras):
            sint = dados.get("sintomas", "")
            if sint:
                nxt = all_paras[i + 1]
                if re.search(r'\[\s+\]', nxt.text):
                    _inline_replace(nxt, sint)

    # ── Assinaturas (tabela 2, sem placeholder, adiciona texto apos label) ────
    sig_c = dados.get("assinatura_cliente", "")
    sig_r = dados.get("assinatura_responsavel", "")
    if sig_c:
        tbl2.rows[0].cells[0].paragraphs[0].add_run(f" {sig_c}")
    if sig_r:
        tbl2.rows[0].cells[1].paragraphs[0].add_run(f" {sig_r}")

    # ── Fotos ao final do documento ───────────────────────────────────────────
    if fotos:
        doc.add_page_break()
        h = doc.add_paragraph("Fotos")
        try:
            h.style = doc.styles["Heading 2"]
        except KeyError:
            pass
        for foto_path in fotos:
            try:
                doc.add_picture(str(foto_path), width=Inches(5.5))
                doc.add_paragraph()
            except Exception:
                pass

    saida = _Path(caminho_saida)
    saida.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(saida))
